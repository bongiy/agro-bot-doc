from __future__ import annotations

import os
import shutil
from typing import Any, Mapping, Iterable

from docx import Document

from template_vars import SUPPORTED_VARS, EMPTY_VALUE
from template_utils import extract_variables
from ftp_utils import download_file_ftp, upload_file_ftp
from docx2pdf import convert

import subprocess
from datetime import datetime
from decimal import Decimal


def docx_to_pdf(docx_path: str, pdf_path: str) -> None:
    try:
        convert(docx_path, pdf_path)
    except Exception:
        libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
        if not libreoffice:
            raise
        subprocess.run([
            libreoffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            os.path.dirname(pdf_path),
            docx_path,
        ], check=True)
        generated = os.path.join(
            os.path.dirname(pdf_path),
            os.path.splitext(os.path.basename(docx_path))[0] + ".pdf",
        )
        os.replace(generated, pdf_path)


def format_area(area: float | int | str | None) -> str:
    if area is None or area == "":
        return EMPTY_VALUE
    try:
        return f"{float(area):.4f}"
    except Exception:
        return str(area)


def format_money(amount: float | Decimal | str | None) -> str:
    if amount is None or amount == "":
        return EMPTY_VALUE
    try:
        value = float(amount)
    except Exception:
        try:
            value = float(str(amount).replace(",", "."))
        except Exception:
            return str(amount)
    parts = f"{value:,.2f}".split(".")
    parts[0] = parts[0].replace(",", " ")
    return f"{parts[0]},{parts[1]} грн"


def format_share(share: float | str | None) -> str:
    if share is None or share == "":
        return EMPTY_VALUE
    return str(share)


def load_template(remote_path: str, local_dir: str = "temp_docs") -> str:
    os.makedirs(local_dir, exist_ok=True)
    local_path = os.path.join(local_dir, os.path.basename(remote_path))
    download_file_ftp(remote_path, local_path)
    return local_path


def read_template_vars(template_path: str) -> set[str]:
    return set(extract_variables(template_path).keys())


def build_context(template_vars: Iterable[str], values: Mapping[str, Any]) -> dict[str, Any]:
    context: dict[str, Any] = {}
    for var in template_vars:
        if var in SUPPORTED_VARS:
            value = values.get(var)
            context[var] = value if value not in (None, "") else EMPTY_VALUE
        else:
            context[var] = EMPTY_VALUE
    return context


def generate_log(template_name: str, total: int, filled: list[str], missing: list[str], unsupported: list[str]) -> str:
    lines = [f"⚠️ У шаблоні {template_name} знайдено {total} змінних"]
    lines.append(f"✅ Заповнено: {len(filled)}")
    if missing or unsupported:
        lines.append("❌ Не заповнено:")
        for var in missing:
            lines.append(var)
        for var in unsupported:
            lines.append(f"{var} — змінна не підтримується")
        lines.append(f"📝 У шаблон буде підставлено: {EMPTY_VALUE}")
    return "\n".join(lines)


def fill_doc(template_path: str, context: Mapping[str, Any], output_docx: str) -> None:
    doc = Document(template_path)
    placeholders = {f"{{{{{k}}}}}": str(v) for k, v in context.items()}

    def replace_in_paragraph(paragraph):
        for ph, val in placeholders.items():
            if ph in paragraph.text:
                paragraph.text = paragraph.text.replace(ph, val)

    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                for ph, val in placeholders.items():
                    if ph in cell.text:
                        cell.text = cell.text.replace(ph, val)

    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for t in doc.tables:
        replace_in_table(t)
    for section in doc.sections:
        hdr = section.header
        for p in hdr.paragraphs:
            replace_in_paragraph(p)
        for t in hdr.tables:
            replace_in_table(t)
        ftr = section.footer
        for p in ftr.paragraphs:
            replace_in_paragraph(p)
        for t in ftr.tables:
            replace_in_table(t)

    doc.save(output_docx)


def generate_contract(
    template_remote_path: str,
    values: Mapping[str, Any],
    payer_name: str,
    contract_number: str,
    year: int,
    *,
    dev: bool = False,
) -> tuple[str, str]:
    template_local = load_template(template_remote_path)
    template_vars = read_template_vars(template_local)
    context = build_context(template_vars, values)

    filled = [v for v in template_vars if context[v] != EMPTY_VALUE]
    missing = [f"{{{{{v}}}}}" for v in template_vars if context[v] == EMPTY_VALUE and v in SUPPORTED_VARS and values.get(v) in (None, "")]
    unsupported = [f"{{{{{v}}}}}" for v in template_vars if v not in SUPPORTED_VARS]

    log = generate_log(os.path.basename(template_remote_path), len(template_vars), filled, missing, unsupported)

    os.makedirs("temp_docs", exist_ok=True)
    filled_docx = os.path.join("temp_docs", "filled_contract.docx")
    fill_doc(template_local, context, filled_docx)
    pdf_local = os.path.join("temp_docs", "contract.pdf")
    docx_to_pdf(filled_docx, pdf_local)
    os.remove(filled_docx)

    filename = f"\u0414\u043e\u0433\u043e\u0432\u0456\u0440_{contract_number}_{payer_name}.pdf"
    remote_dir = f"contracts/{year}/{payer_name}"
    remote_path = f"{remote_dir}/{filename}"

    if dev:
        os.makedirs(remote_dir, exist_ok=True)
        shutil.move(pdf_local, remote_path)
    else:
        upload_file_ftp(pdf_local, remote_path)
        os.remove(pdf_local)

    os.remove(template_local)

    return remote_path, log



async def generate_contract_v2(contract_id: int) -> tuple[str, str]:
    """Generate a contract PDF for the given contract ID.

    The contract along with its related payer, company and land plots are
    loaded from the database. The first active agreement template is used
    to produce the PDF. The resulting file is uploaded via FTP and a record
    is stored in ``uploaded_docs``.

    Parameters
    ----------
    contract_id: int
        Identifier of the contract to generate.

    Returns
    -------
    tuple[str, str]
        Remote path to the uploaded PDF and generation log.
    """
    import sqlalchemy
    from db import (
        database,
        Contract,
        Company,
        Payer,
        LandPlot,
        LandPlotOwner,
        ContractLandPlot,
        UploadedDocs,
        get_agreement_templates,
    )
    from template_utils import analyze_template

    contract = await database.fetch_one(
        sqlalchemy.select(Contract).where(Contract.c.id == contract_id)
    )
    if not contract:
        raise ValueError("Contract not found")

    if not contract["company_id"]:
        raise RuntimeError("Contract has no company specified")
    if not contract["payer_id"]:
        raise RuntimeError("Contract has no payer specified")

    company = await database.fetch_one(
        sqlalchemy.select(Company).where(Company.c.id == contract["company_id"])
    )
    payer = await database.fetch_one(
        sqlalchemy.select(Payer).where(Payer.c.id == contract["payer_id"])
    )
    if not company or not payer:
        raise RuntimeError("Required contract relations not found")

    lands = await database.fetch_all(
        sqlalchemy.select(LandPlot)
        .join(ContractLandPlot, LandPlot.c.id == ContractLandPlot.c.land_plot_id)
        .where(ContractLandPlot.c.contract_id == contract_id)
    )
    if not lands:
        raise RuntimeError("Contract has no land plots")

    land_ids = [l["id"] for l in lands]
    owners_rows = await database.fetch_all(
        sqlalchemy.select(
            LandPlotOwner.c.land_plot_id,
            LandPlotOwner.c.payer_id,
            LandPlotOwner.c.share,
            Payer.c.name,
        )
        .join(Payer, Payer.c.id == LandPlotOwner.c.payer_id)
        .where(LandPlotOwner.c.land_plot_id.in_(land_ids))
    )

    land_owner_map: dict[int, dict[int, float]] = {}
    payer_shares: dict[int, dict[str, Any]] = {}
    for r in owners_rows:
        land_owner_map.setdefault(r["land_plot_id"], {})[r["payer_id"]] = r["share"]
        info = payer_shares.setdefault(
            r["payer_id"], {"name": r["name"], "share": 0.0}
        )
        info["share"] += float(r["share"])

    payer_share_value = payer_shares.get(contract["payer_id"], {}).get("share")
    payers_list = "\n".join(
        f"{info['name']} — {format_share(info['share'])}" for info in payer_shares.values()
    )
    plots_table = "\n".join(
        f"{l['cadaster']} — {format_area(l['area'])}" for l in lands
    )
    first_land = lands[0]
    first_land_share = land_owner_map.get(first_land["id"], {}).get(contract["payer_id"])

    templates = await get_agreement_templates(True)
    if not templates:
        raise RuntimeError("No active agreement template")
    template = templates[0]

    tmp_doc = f"temp_docs/template_{contract_id}.docx"
    download_file_ftp(template["file_path"], tmp_doc)

    variables = {
        "contract_number": contract["number"],
        "contract_date_signed": contract["date_signed"].strftime("%d.%m.%Y"),
        "contract_date_from": contract["date_valid_from"].strftime("%d.%m.%Y"),
        "contract_date_to": contract["date_valid_to"].strftime("%d.%m.%Y"),
        "contract_term": contract["duration_years"],
        "contract_rent": float(contract["rent_amount"]),
        "company_name": company["full_name"],
        "company_code": company["edrpou"],
        "company_address": company["address_legal"],
        "company_director": company["director"],
        "payer_name": payer["name"],
        "payer_tax_id": payer["ipn"],
        "payer_birthdate": payer["birth_date"],
        "payer_passport": (
            f"{(payer['passport_series'] or payer['id_number'] or '')} "
            f"{payer['passport_number'] or ''} "
            f"{payer['passport_issuer'] or payer['idcard_issuer'] or ''} "
            f"{payer['passport_date'] or payer['idcard_date'] or ''}"
        ).strip(),
        "payer_address": (
            f"{payer['oblast']} обл., {payer['rayon']} р-н, с. {payer['selo']}, "
            f"вул. {payer['vul']}, буд. {payer['bud']}"
            f"{(', кв. ' + payer['kv']) if payer['kv'] else ''}"
        ),
        "payer_share": format_share(payer_share_value),
        "payer_phone": payer["phone"],
        "payer_bank_card": payer["bank_card"],
        "payers_list": payers_list,
        "plot_cadastre": first_land["cadaster"],
        "plot_area": format_area(first_land["area"]),
        "plot_share": format_share(first_land_share),
        "plot_ngo": format_money(first_land["ngo"]),
        "plot_location": ", ".join(
            filter(None, [first_land["council"], first_land["district"], first_land["region"]])
        ),
        "plots_table": plots_table,
        "today": datetime.utcnow().strftime("%d.%m.%Y"),
        "year": datetime.utcnow().year,
    }

    # Analyze template for unresolved placeholders (result not used yet)
    analyze_template(tmp_doc, variables)
    if os.path.exists(tmp_doc):
        os.remove(tmp_doc)

    remote_path, gen_log = generate_contract(
        template["file_path"],
        variables,
        payer_name=payer["name"],
        contract_number=contract["number"],
        year=datetime.utcnow().year,
    )

    await database.execute(
        UploadedDocs.delete().where(
            (UploadedDocs.c.entity_type == "contract")
            & (UploadedDocs.c.entity_id == contract_id)
            & (UploadedDocs.c.doc_type == "generated")
        )
    )
    await database.execute(
        UploadedDocs.insert().values(
            entity_type="contract",
            entity_id=contract_id,
            doc_type="generated",
            remote_path=remote_path,
        )
    )

    return remote_path, gen_log

